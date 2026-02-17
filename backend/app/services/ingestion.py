from fastapi import UploadFile, HTTPException
import shutil
import os
from typing import List, Optional
from uuid import uuid4
from sqlalchemy.orm import Session
from app.models.sql_models import Document, DocumentChunk, ContentType
from app.services.ai.factory import AIFactory
import logging
try:
    import pypdf
except ImportError:
    pypdf = None
try:
    import docx
except ImportError:
    docx = None
from langchain.text_splitter import RecursiveCharacterTextSplitter

class IngestionService:
    UPLOAD_DIR = "uploads"
    
    def __init__(self, db: Session):
        self.db = db
        self.ai_service = AIFactory.get_service()
        if not os.path.exists(self.UPLOAD_DIR):
            os.makedirs(self.UPLOAD_DIR)

    def _extract_text_from_pdf(self, file_path: str) -> str:
        if not pypdf:
            return "[PDF Support Missing]"
        text = ""
        try:
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        except Exception as e:
            logging.error(f"PDF Error: {e}")
            return f"[PDF Error: {e}]"
        return text

    def _extract_text_from_docx(self, file_path: str) -> str:
        if not docx:
             return "[DOCX Support Missing]"
        try:
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            logging.error(f"DOCX Error: {e}")
            return f"[DOCX Error: {e}]"

    async def _extract_text(self, file_path: str, file_ext: str) -> str:
        if file_ext == "pdf":
            return self._extract_text_from_pdf(file_path)
        elif file_ext in ["docx", "doc"]:
            return self._extract_text_from_docx(file_path)
        elif file_ext == "txt":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            except Exception as e:
                return f"[TXT Error: {e}]"
        else:
            return f"[Unsupported Format: {file_ext}]"

    async def process_upload(self, file: UploadFile, user_id: str) -> Document:
        # 1. Save File Locally
        file_ext = file.filename.split(".")[-1].lower()
        file_id = str(uuid4())
        safe_filename = f"{file_id}.{file_ext}"
        file_path = os.path.join(self.UPLOAD_DIR, safe_filename)
        
        try:
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
        except Exception as e:
            logging.error(f"File save failed: {e}")
            raise HTTPException(status_code=500, detail="Could not save file")

        # 2. Extract Text
        content_text = await self._extract_text(file_path, file_ext)
        
        # 3. Create Document Record
        doc = Document(
            id=file_id,
            user_id=user_id,
            title=file.filename,
            content=content_text,
            file_path=file_path,
            type=ContentType.TEXT # Simplify to TEXT for now
        )
        self.db.add(doc)
        # We don't commit yet, wait for chunks

        # 4. Chunking & Vectorization
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        chunks = text_splitter.split_text(content_text if content_text else "Empty content")
        
        # Process chunks
        for i, chunk_text in enumerate(chunks):
            # Skip empty chunks
            if not chunk_text.strip():
                continue
                
            try:
                embedding = await self.ai_service.get_embeddings(chunk_text)
                
                chunk_record = DocumentChunk(
                    id=str(uuid4()),
                    document_id=file_id,
                    content=chunk_text,
                    embedding=embedding,
                    chunk_index=i,
                    meta_data={"source": file.filename, "chunk": i}
                )
                self.db.add(chunk_record)
            except Exception as e:
                logging.error(f"Chunk vectorization failed for {i}: {e}")
                # Continue processing others? or fail?
                # For now, continue but log error
                pass

        self.db.commit()
        self.db.refresh(doc)
        
        return doc
